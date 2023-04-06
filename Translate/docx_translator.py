from io import BytesIO
from docx import Document

def transform_paragraphs(paragraphs, translator, dest_lang):
    '''
    Only support .docx type file of Microsoft Office Word
    Args:
        paragraphs : Paragraph Obj
        translator : Translator Obj
        dest_lang : target language to translate
    '''
    if dest_lang.lower() == 'english':
        src_lang, tgt_lang = 'ko_KR', 'en_XX'
    elif dest_lang.lower() == 'korean':
        src_lang, tgt_lang = 'en_XX', 'ko_KR'
    else:
        raise ValueError("Only support languages are korean and english")

    for paragraph in paragraphs:
        texts = [run.text for run in paragraph.runs]
        if texts != []:
            for i, run in enumerate(paragraph.runs):
                if i == 0:
                    if run.text == '':
                        continue
                    else:
                        translated = translator.generate([' '.join(texts)], src_lang=src_lang, tgt_lang=tgt_lang)
                        run.text = '\n'.join([s['translated'] for s in translated])
                else:
                    run.clear()
        else:
            continue

def trans_with_preserving_form(uploaded_file, translator, dest_lang='Enlgish'):
    """
    Function to display the translated Word document
    Args:
        uploaded_file : 
    """
    docx = Document(BytesIO(uploaded_file))
    
    #Header Part
    print("Translating Header Part")
    for section in docx.sections:
        #header contents
        transform_paragraphs(section.header.paragraphs, translator=translator, dest_lang=dest_lang)
        #header tables
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    transform_paragraphs(cell.paragraphs, translator=translator, dest_lang=dest_lang)

    #Body Part
    print("Translating Body Part")
    #body contents
    transform_paragraphs(docx.paragraphs, translator=translator, dest_lang=dest_lang)
    #body tables
    for table in docx.tables:
        for row in table.rows:
            for cell in row.cells:
                transform_paragraphs(cell.paragraphs, translator=translator, dest_lang=dest_lang)

    #Footer Part
    print("Translating Footer Part")
    for section in docx.sections:
        #footer contents
        transform_paragraphs(section.footer.paragraphs, translator=translator, dest_lang=dest_lang)
        #footer tables
        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    transform_paragraphs(cell.paragraphs, translator=translator, dest_lang=dest_lang)

    return docx
